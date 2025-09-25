import os
from typing import Dict, Any
from jinja2 import Environment, FileSystemLoader, select_autoescape
from docx import Document

TEMPL_DIR = os.path.join(os.path.dirname(__file__), "templates")
OUT_DIR = os.path.join(os.path.dirname(__file__), "out")

def render_markdown(context: Dict[str, Any]) -> str:
    env = Environment(
        loader=FileSystemLoader(TEMPL_DIR),
        autoescape=select_autoescape()
    )
    tmpl = env.get_template("proposal.md.j2")
    return tmpl.render(**context)

def save_docx_from_context(context: Dict[str, Any], path: str) -> None:
    # A lightweight DOCX writer (doesn't fully support markdown)
    doc = Document()
    doc.add_heading(context.get("title","Sales Proposal"), level=0)

    # Sections
    def add_heading(doc, text): doc.add_heading(text, level=1)
    def add_par(doc, text): 
        for line in text.splitlines():
            doc.add_paragraph(line)

    add_heading(doc, "Executive Summary")
    add_par(doc, context.get("solution_overview",""))

    add_heading(doc, "Customer Objectives")
    for o in context.get("objectives", []):
        doc.add_paragraph(f"• {o}")

    add_heading(doc, "Proposed Solution")
    for s in context.get("selected_skus", []):
        doc.add_paragraph(f"- {s['sku']}: {s.get('name','')} — {s.get('reason','')}")

    add_heading(doc, "Architecture")
    for a in context.get("architecture", []):
        doc.add_paragraph(f"- {a['component']}: {a['role']}")

    add_heading(doc, "Implementation Plan")
    for p in context.get("implementation_plan", []):
        doc.add_paragraph(f"- {p.get('phase','')}: {p.get('weeks','')} weeks")

    add_heading(doc, "Commercials")
    pr = context.get("pricing_breakdown", {})
    items = pr.get("items", [])
    for it in items:
        doc.add_paragraph(f"- {it['name']} (SKU {it['sku']}) x{it['qty']}: {pr.get('currency','')} {it['subtotal']:.2f}")
    doc.add_paragraph(f"Subtotal: {pr.get('currency','')} {pr.get('subtotal',0):.2f}")
    doc.add_paragraph(f"Discount ({pr.get('discount_pct',0)}%): -{pr.get('currency','')} {pr.get('discount_amount',0):.2f}")
    doc.add_paragraph(f"Tax ({pr.get('tax_pct',0)}%): {pr.get('currency','')} {pr.get('tax_amount',0):.2f}")
    doc.add_paragraph(f"Total: {pr.get('currency','')} {pr.get('total',0):.2f}")

    add_heading(doc, "Risks & Mitigations")
    for rm in context.get("risks_mitigations", []):
        doc.add_paragraph(f"- Risk: {rm.get('risk','')} | Mitigation: {rm.get('mitigation','')}" )

    add_heading(doc, "Assumptions & Terms")
    for a in context.get("assumptions", []):
        doc.add_paragraph(f"- {a}")

    add_heading(doc, "Citations")
    for c in context.get("citations", []):
        doc.add_paragraph(c)

    doc.save(path)
